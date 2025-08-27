import time
from docx import Document
from docx.shared import Inches
from selenium import webdriver
import os

def before_all(context):
    context.driver = webdriver.Edge()
    context.driver.maximize_window()
    context.doc = Document()
    context.screenshots_path = f"{os.getcwd()}/reports"


def before_step(context, step):
    filename = f"{context.screenshots_path}/{step.name}.png"
    context.driver.save_screenshot(filename)
    context.doc.add_paragraph(step.name)
    context.doc.add_picture(filename, width=Inches(4.0))

def after_all(context):
    context.doc.save(f"{context.screenshots_path}/report.docx")
    context.driver.quit()